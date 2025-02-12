import requests
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswAlgorithmConfiguration,
    VectorSearch,
    VectorSearchAlgorithmConfiguration,
    VectorSearchProfile,
)
from azure.search.documents.indexes.models import (
    SearchIndex,
    SimpleField,
    SearchableField,
    SearchField
)
import pdfplumber
from docx import Document
import io
import openpyxl
import os
from azure.storage.blob import BlobServiceClient
import logging
# SharePoint and Azure configurations
SHAREPOINT_DRIVE_ID = ""
ACCESS_TOKEN = ""
FOLDER_ID = ""
CONNECTION_STRING = ""
CONTAINER_NAME = "acg-sample-images"

def get_sharepoint_items(access_token, folder_id=None):
    if folder_id:
        url = f"https://graph.microsoft.com/v1.0/drives/{SHAREPOINT_DRIVE_ID}/items/{folder_id}/children"
    else:
        url = f"https://graph.microsoft.com/v1.0/drives/{SHAREPOINT_DRIVE_ID}/root/children"

    headers = {
        "Authorization": "Bearer " + access_token,
        "Accept": "application/json"
    }
    response = requests.get(url, headers=headers)
    data = response.json()
    return data['value']

def download_sharepoint_file(file_id, access_token):
    content_url = f"https://graph.microsoft.com/v1.0/drives/{SHAREPOINT_DRIVE_ID}/items/{file_id}/content"
    headers = {
        "Authorization": "Bearer " + access_token,
    }
    response = requests.get(content_url, headers=headers)
    if response.status_code == 200:
        return response.content
    else:
        print(f"Failed to download file with ID '{file_id}'")
        return None

def get_permissions(file_id, access_token):
    permissions_url = f"https://graph.microsoft.com/v1.0/drives/{SHAREPOINT_DRIVE_ID}/items/{file_id}/permissions"
    headers = {
        "Authorization": "Bearer " + access_token,
        "Accept": "application/json"
    }
    response = requests.get(permissions_url, headers=headers)
    permissions = []
    if response.status_code == 200:
        permissions_data = response.json().get("value", [])
        for permission_data in permissions_data:
            granted_to = permission_data.get("grantedTo", None)
            if granted_to:
                if "user" in granted_to:
                    user_or_group = granted_to["user"].get("displayName", "N/A")
                elif "group" in granted_to:
                    user_or_group = granted_to["group"].get("displayName", "N/A")
                if user_or_group != "N/A":
                    permissions.append(f" {user_or_group}")
    return permissions

def extract_text_from_pdf(file_content):
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def process_docx(docx_content):
        """Extract content under each heading from a DOCX file, without printing the headings."""
        try:
            doc = Document(io.BytesIO(docx_content))
            heading_content = {}
            current_heading = None
            current_text = []

            for paragraph in doc.paragraphs:
    
                if paragraph.style.name.startswith('Heading'):
                    if current_heading:  
                        heading_content[current_heading] = "\n".join(current_text).strip()
                    current_heading = paragraph.text 
                    current_text = []  
                else:
                    current_text.append(paragraph.text) 

            
            if current_heading:
                heading_content[current_heading] = "\n".join(current_text).strip()

        except Exception as e:
            logging.error(f"Error processing DOCX file: {str(e)}")
            return {}
        
        return heading_content  

def extract_text_from_excel(file_content):
    file_str = ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_content))
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        file_str += str(cell.value) + " "
    except Exception as e:
        file_str = f"Error extracting text from Excel: {e}"
    return file_str

def extract_text_from_file(file_content, file_extension):
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_extension in ['docx', 'doc']:
        return process_docx(file_content)
    elif file_extension in ['xlsx', 'xls']:
        return extract_text_from_excel(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def upload_image_to_blob(image_data, container_name, blob_name, connection_string):
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_name)

    # Upload the image data
    blob_client.upload_blob(image_data, overwrite=True)
    print(f"Uploaded {blob_name} to Azure Blob Storage.")

    # Construct the Blob URL
    return blob_client.url  # Return the URL of the uploaded blob

def extract_images_from_docx(file_content, file_name, container_name, connection_string):
    doc = Document(io.BytesIO(file_content))
    image_urls = []  # List to hold paths of extracted images
    img_counter = {}

    current_heading = "no_heading"
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            current_heading = para.text.strip().replace(' ', '_')
            img_counter[current_heading] = 1  # Initialize counter for each heading

        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                blip = run._element.xpath('.//a:blip')[0]
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                image_part = doc.part.related_parts[rId]
                image_data = image_part.blob

                # Construct the image filename using the format: file_name_current_heading_imgcounter
                sanitized_file_name = file_name.replace(" ", "_").replace(".docx", "")  # Sanitize the file name
                filename = f"{sanitized_file_name}_{current_heading}_img_{img_counter[current_heading]}.png"

                # Save image to Azure Blob and get the URL
                image_url = upload_image_to_blob(image_data, container_name, filename, connection_string)

                img_counter[current_heading] += 1
                image_urls.append((current_heading, image_url))  # Store heading and image URL

    print(f"Image extraction complete!")
    return image_urls

# Usage example
if __name__ == "__main__":
    # Get SharePoint files
    files = get_sharepoint_items(ACCESS_TOKEN, folder_id=FOLDER_ID)
    
    for file in files:
        file_id = file["id"]
        file_name = file["name"]
        file_content = download_sharepoint_file(file_id, ACCESS_TOKEN)
        
        if file_content:
            # Extract text based on file type
            if file_name.endswith(".pdf"):
                extracted_text = extract_text_from_pdf(file_content)
            elif file_name.endswith(".docx"):
                extracted_text = process_docx(file_content)
                # Extract images if it's a Word document
                images = extract_images_from_docx(file_content, file_name,CONTAINER_NAME, CONNECTION_STRING)
                print(f"Extracted images from {file_name}: {images}")
            elif file_name.endswith(".xlsx"):
                extracted_text = extract_text_from_excel(file_content)
            
            print(f"Extracted text from {file_name}:")
            print(extracted_text)

    